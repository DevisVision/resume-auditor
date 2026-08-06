from __future__ import annotations

import io
import fitz

from pypdf import PdfReader
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def extract_text(file_name: str, file_bytes: bytes) -> str:
    """Extract plain text AND embedded hyperlink URIs from an uploaded resume file."""
    name = file_name.lower().strip()
    if name.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if name.endswith(".docx"):
        return _extract_docx(file_bytes)
    if name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {file_name}. Use PDF, DOCX or TXT.")


def _extract_pdf(data: bytes) -> str:
    """
    Extract text from PDF.

    Strategy:
    1. Try pypdf.
    2. If very little text is extracted, fall back to PyMuPDF.
    3. Preserve embedded hyperlinks.
    """

    parts = []
    links = []

    # -------------------------------------------------
    # Primary Extraction using pypdf
    # -------------------------------------------------
    try:
        reader = PdfReader(io.BytesIO(data))

        for page in reader.pages:
            try:
                txt = page.extract_text()
                if txt:
                    parts.append(txt)
            except Exception:
                continue

        # Extract hyperlinks
        for page in reader.pages:
            try:
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        obj = annot.get_object()

                        if obj.get("/Subtype") == "/Link":
                            action = obj.get("/A")

                            if action and "/URI" in action:
                                links.append(str(action["/URI"]))

            except Exception:
                continue

    except Exception:
        pass

    full_text = "\n".join(parts).strip()

    # -------------------------------------------------
    # Fallback using PyMuPDF
    # -------------------------------------------------
    if len(full_text) < 300:

        try:
            doc = fitz.open(stream=data, filetype="pdf")

            parts = []

            for page in doc:
                try:
                    txt = page.get_text("text")

                    if txt:
                        parts.append(txt)

                except Exception:
                    continue

            full_text = "\n".join(parts).strip()

        except Exception:
            pass

    # -------------------------------------------------
    # Append hyperlinks
    # -------------------------------------------------
    if links:

        full_text += (
            "\n\n--- EXTRACTED EMBEDDED LINKS ---\n"
            + "\n".join(sorted(set(links)))
        )

    return full_text


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
                    
    # Extract embedded DOCX hyperlinks from back-end document mapping packages
    links = []
    try:
        for rel_id, rel in doc.part.rels.items():
            if rel.reltype == RT.HYPERLINK:
                links.append(rel.target_ref)
    except Exception:
        pass

    full_text = "\n".join(parts).strip()
    if links:
        full_text += "\n\n--- EXTRACTED EMBEDDED LINKS ---\n" + "\n".join(set(links))
        
    return full_text